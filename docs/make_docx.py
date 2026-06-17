#!/usr/bin/env python3
# Сборка Word-документа исследования Uni-one (текст + таблицы + диаграммы).
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

HERE = os.path.dirname(__file__)
CH = os.path.join(HERE, "charts")

ORANGE = RGBColor(0xF5, 0x51, 0x2F)
BLUE = RGBColor(0x2F, 0x6B, 0xFF)
INK = RGBColor(0x1C, 0x22, 0x30)
GREY = RGBColor(0x6B, 0x72, 0x80)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.font.color.rgb = INK


def h(text, level=1, color=BLUE):
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    run.font.color.rgb = color
    return p


def para(text, italic=False, size=11, color=INK, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hd in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        run = c.paragraphs[0].add_run(hd)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10.5)
        # фон ячейки
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        sh = OxmlElement("w:shd")
        sh.set(qn("w:fill"), "2F6BFF")
        c._tc.get_or_add_tcPr().append(sh)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    return t


def chart(name, width=6.2):
    path = os.path.join(CH, name)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


# ===== Титул =====
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Uni·one")
r.bold = True
r.font.size = Pt(34)
r.font.color.rgb = BLUE
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Исследование пользователей · Customer Development")
r.font.size = Pt(14)
r.font.color.rgb = GREY
para("Приложение для знакомств и социализации студентов",
     italic=True, color=GREY).alignment = WD_ALIGN_PARAGRAPH.CENTER
note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = note.add_run("⚠ Данные модельные (иллюстративные) — оформлены для презентации")
r.italic = True
r.font.size = Pt(9.5)
r.font.color.rgb = ORANGE

# ===== Метод =====
h("1. Методология", 1)
para("Выборка: 45 человек — студенты 1 курса бакалавриата / абитуриенты "
     "и 1 курса магистратуры.")
for s in [
    "Проблемное интервью (адаптация и поиск друзей) — до показа продукта.",
    "Вопросы на потребность в подобном приложении.",
    "Тест прототипа Uni-one (свободное использование) → сбор проблем.",
    "Синтез инсайтов → доработки продукта.",
]:
    doc.add_paragraph(s, style="List Number")

# ===== Демография =====
h("2. Демография респондентов", 1)
table(["Параметр", "Распределение"],
      [["Ступень", "1 курс бак./абитуриенты — 29 (64%) · 1 курс магистратуры — 16 (36%)"],
       ["Пол", "Женский — 24 (53%) · Мужской — 21 (47%)"],
       ["Откуда", "Иногородние — 27 (60%) · Москва — 18 (40%)"]],
      widths=[1.6, 4.8])
doc.add_paragraph()
chart("01_respondents_stage.png", 3.6)
chart("02_respondents_origin.png", 3.6)

# ===== Блок 1 =====
doc.add_page_break()
h("3. Блок 1. Адаптация и поиск друзей", 1)
para("10 вопросов до показа продукта.", italic=True, color=GREY)
b1 = [
    ["1", "Насколько легко находить друзей?", "Очень сложно 38% · скорее 33% · нейтр. 16% · легко 13% → 71% сложно"],
    ["2", "Близких друзей за 1-й месяц", "0–1: 49% · 2–3: 31% · 4+: 20%"],
    ["3", "Главный барьер знакомства", "Стеснение 42% · нет повода 27% · мало времени 18% · компании 13%"],
    ["4", "Одиночество в первые недели", "Часто 29% · иногда 47% · нет 24% → 76%"],
    ["5", "Откуда узнаёшь о событиях", "Чаты/ТГ 40% · друзья 27% · никак 18% · стенды 15%"],
    ["6", "Инфо в одном месте найти сложно?", "Да 69%"],
    ["7", "Ходишь на события один?", "«Не с кем» 38% · некомфортно 29% · спокойно 33%"],
    ["8", "Что мешает ходить", "Не с кем 44% · не знаю что есть 31% · стесняюсь 25%"],
    ["9", "Знакомиться онлайн заранее?", "Да 73%"],
    ["10", "Круг общения влияет на мотивацию?", "Сильно 58% · немного 31% · нет 11%"],
]
table(["#", "Вопрос", "Ответы (n=45)"], b1, widths=[0.4, 2.6, 3.6])
para("Вывод: проблема не «не с кем общаться», а нет повода и безопасной "
     "среды + информация разрознена.", bold=True)
doc.add_paragraph()
chart("03_friends_difficulty.png", 5.6)
chart("04_barriers.png", 5.6)
chart("05_info_sources.png", 5.6)

# ===== Блок 2 =====
doc.add_page_break()
h("4. Блок 2. Потребность в приложении", 1)
para("10 вопросов после презентации идеи.", italic=True, color=GREY)
b2 = [
    ["1", "Пользуешься дейтинг/нетворкингом?", "Да 51% · нет 49%"],
    ["2", "Хотел бы приложение для своего вуза?", "Да 80%"],
    ["3", "Что важнее — знакомства или события?", "Связка обоих 56% · события 24% · знакомства 20%"],
    ["4", "Идея «люди через общие события»?", "Очень 47% · скорее да 31% → 4.3/5"],
    ["5", "Фича «позвать вместе на событие»?", "Да 71%"],
    ["6", "Групповой чат события нужен?", "Да 67%"],
    ["7", "Важно видеть, куда идут друзья?", "Да 62%"],
    ["8", "Делить контакты на круги?", "Интересно 54% · нейтр. 30% · не нужно 16%"],
    ["9", "XP/ачивки мотивируют?", "Да 49% · возможно 33% · нет 18%"],
    ["10", "Установил бы приложение?", "Точно да 38% · скорее да 40% · вряд ли 22% → 78%"],
]
table(["#", "Вопрос", "Ответы (n=45)"], b2, widths=[0.4, 2.6, 3.6])
para("Вывод: спрос есть, ключевая ценность — симбиоз знакомств и событий, "
     "а не «ещё один Tinder».", bold=True)
doc.add_paragraph()
chart("07_install_intent.png", 6.0)
chart("06_key_metrics.png", 6.2)

# ===== Блок 3 =====
doc.add_page_break()
h("5. Тест прототипа: что нашли", 1)
b3 = [
    ["Звёздочка непонятна — думали «суперлайк»", "16 (36%)"],
    ["Неясно, с чего начать / что умеет", "14 (31%)"],
    ["Хотели менять интересы после онбординга", "12 (27%)"],
    ["Хотели включать несколько фильтров сразу", "10 (22%)"],
    ["Не видно, кто из друзей идёт", "9 (20%)"],
    ["Вкладка «Пары» лишняя/пустоватая", "8 (18%)"],
    ["В свайпе не хватало «общего» с человеком", "8 (18%)"],
    ["Хотели ходить компанией, а не 1-на-1", "7 (16%)"],
]
table(["Находка", "Частота"], b3, widths=[5.0, 1.4])
doc.add_paragraph()
chart("08_prototype_findings.png", 6.2)

# ===== Инсайты =====
doc.add_page_break()
h("6. Инсайты → доработки продукта", 1)
ins = [
    ["Звезда дублировала лайк и путала", "⭐ Фича «Позвать на мероприятие» (свайп + чат)"],
    ["Барьер №1 — «не с кем пойти»", "👥 «Пойти компанией» + групповой чат события"],
    ["Хотят видеть друзей в контексте событий", "👀 «Где друзья сегодня» + друзья на карточках"],
    ["Нужно «общее» для старта разговора", "🔥 Бейдж «Вы оба идёте на…» в свайпах"],
    ["Непонятно, с чего начать", "🧭 Онбординг-тур по возможностям"],
    ["Хотят управлять профилем", "✏️ Редактирование интересов"],
    ["Фильтры неудобны", "✅ Мультивыбор фильтров"],
    ["«Пары» лишние", "💬 Убрали вкладку, мэтчи → сторис в чатах"],
    ["Нужна вовлечённость", "🏆 Геймификация (XP, бейджи, статусы)"],
]
table(["Инсайт из исследования", "Что изменили в продукте"], ins, widths=[3.0, 3.4])
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("Главный инсайт: ")
r.bold = True
r.font.color.rgb = ORANGE
r2 = p.add_run("студенты готовы знакомиться, если есть повод и безопасная "
               "среда — поэтому продукт строится как симбиоз знакомств и "
               "мероприятий, а не как отдельный дейтинг или афиша.")
r2.bold = True

out = os.path.join(HERE, "Uni-one_research.docx")
doc.save(out)
print("Сохранено:", out)
